from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.document import Document as DocumentObject
import unified_statistical_analysis_file

# Count the number of all paragraph elements in the document
def count_total_paragraphs(document: DocumentObject) -> int:
    count = 0
    for _ in document.paragraphs:
        count += 1
    return count

# Count the number of all run elements in the document
def count_total_runs_elements(document: DocumentObject) -> int:
    count = 0
    for paragraph in document.paragraphs:
        for _ in paragraph.runs:
            count += 1
    return count

# Count the number of all text elements in the document
def count_total_text_elements(document: DocumentObject) -> int:
    # Same principle as document.part.element.xpath("//w:document/w:body/w:p/w:r/w:t")
    text_element_count = 0
    text_element = f".//{qn('w:t')}"
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text_element_in_run_count = len(run._r.findall(text_element)) # ElementPath findall() function - returns a list of matching Elements
            text_element_count += text_element_in_run_count
    return text_element_count

# Count text element number in a single paragraph element
def count_text_elements_per_paragraph(document: DocumentObject) -> list[int]:
    text_element_per_paragraph = []
    text_element = f".//{qn('w:t')}"
    for paragraph in document.paragraphs:
        text_element_in_run_count = len(paragraph._p.findall(text_element))
        text_element_per_paragraph.append(text_element_in_run_count)
    return text_element_per_paragraph

# Count text element number in a single run element
def count_text_elements_per_run(document: DocumentObject) -> list[int]:
    text_element_per_run = []
    text_element = f".//{qn('w:t')}"
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text_element_in_run_count = len(run._r.findall(text_element))
            text_element_per_run.append(text_element_in_run_count)
    return text_element_per_run

### Main function ###
def element_count_analysis(path: Path, data_set: str, chosen_file: bool) -> tuple[list[list], int]:
    document = Document(str(path))
    total_paragraph_count = count_total_paragraphs(document)
    total_run_element_count = count_total_runs_elements(document)
    total_text_elements = count_total_text_elements(document)
    text_element_per_paragraph_list = count_text_elements_per_paragraph(document)
    text_element_per_run_list = count_text_elements_per_run(document)

    # Data export
    if not chosen_file:
        data_to_csv = [
            ["Data set", data_set],
            ["Document Name", "Total Paragraph Count", "Total Run Element Count", "Total Text Element Count"],
            [path.stem, total_paragraph_count, total_run_element_count, total_text_elements]
        ]
    else:
        data_to_csv = [
            ["Document Name", path.stem],
            ["Data set", data_set],
            ["Total Paragraph Count", total_paragraph_count],
            ["Total Run Element Count", total_run_element_count],
            ["Total Text Element Count", total_text_elements],
            ["Text Elements Per Paragraph", *text_element_per_paragraph_list],
            ["Text Elements Per Run", *text_element_per_run_list]
        ]
    return data_to_csv, 2

def main() -> None:
    unified_statistical_analysis_file.statistical_analysis('1_element_count', 'TEST_0', element_count_analysis)
    unified_statistical_analysis_file.statistical_analysis('1_element_count', None, element_count_analysis)

if __name__ == "__main__":
    main()