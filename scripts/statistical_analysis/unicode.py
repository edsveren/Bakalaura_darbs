import re
from docx import Document
from pathlib import Path

def count_text_elements(document) -> int:
    root = document.part.element
    docTextElement = root.xpath("//w:document/w:body/w:p/w:r/w:t")
    return len(docTextElement)

def count_chars_in_paragraphs(document) -> int:
    char_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
        chars = re.findall(r'[\s\S]', text, flags=re.UNICODE)
        char_count += len(chars)
    return char_count

def count_non_ascii_chars_in_paragraphs(document) -> int:
    char_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.replace('\xa0', '\x20')
        chars = re.findall(r'[^\x00-\x7F]', text, flags=re.UNICODE)
        char_count += len(chars)
    return char_count

docPath_0 = Path("data_set/clean_files/TEST_0.docx")
docPath_1 = Path("data_set/stego-files/stego-method_1/TEST_0.docx")
docPath_2 = Path("data_set/stego-files/stego-method_4/TEST_0.docx")
docPath_3 = Path("data_set/stego-files/stego-method_5/TEST_0.docx")
docPath_4= Path("data_set/stego-files/stego-method_6/TEST_0.docx")

paths = [docPath_0, docPath_1, docPath_2, docPath_3, docPath_4]

for path in paths:
    print("")
    print(f"DOCUMENT: {path}")
    document = Document(path)
    char_count = count_chars_in_paragraphs(document)
    non_ascii_char_count = count_non_ascii_chars_in_paragraphs(document)
    char_to_non_ascii_char_ratio = 100 * (non_ascii_char_count/char_count)
    print(f"Non-ASCII char count: {non_ascii_char_count}")
    print(f"Char count: {char_count}")
    print(f"Char count to Non-ASCII char count ratio: {char_to_non_ascii_char_ratio}")