import re
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentObject
import scripts.statistical_analysis.unified_statistical_analysis_file as unified_statistical_analysis_file

# Count the amount of individual characters in each paragraph element
def count_chars_in_each_paragraph(document: DocumentObject) -> int:
    char_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.replace('\xa0', '\x20')
        chars = re.findall(r'[\s\S]', text, flags=re.UNICODE)
        char_count += len(chars)
    return char_count

# Count the amount of individual non-ASCII characters in each paragraph element
def count_non_ascii_chars_in_each_paragraph(document: DocumentObject) -> tuple[int, list[int]]:
    char_count_total = 0
    char_count_paragraph = []
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.replace('\xa0', '\x20')
        paragraph_chars = re.findall(r'[^\x00-\x7F]', paragraph_text, flags=re.UNICODE)
        paragraph_chars_count = len(paragraph_chars)
        char_count_paragraph.append(paragraph_chars_count)
        char_count_total += paragraph_chars_count
    return char_count_total, char_count_paragraph

# Count the amount of individual non-ASCII characters in each run element
def count_non_ascii_chars_in_runs(document: DocumentObject) -> list[int]:
    char_count_run = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run_text = run.text.replace('\xa0', '\x20')
            run_chars = re.findall(r'[^\x00-\x7F]', run_text, flags=re.UNICODE)
            char_count_run.append(len(run_chars))
    return char_count_run

### Main function ###
def unicode_analysis(path: Path, data_set: str, chosen_file: bool) -> tuple[list[list], int]:
    document = Document(str(path))
    total_char_count = count_chars_in_each_paragraph(document)
    total_non_ascii_char_count, non_ascii_char_count_paragraphs = count_non_ascii_chars_in_each_paragraph(document)
    if total_char_count != 0:
        total_char_to_non_ascii_char_ratio = round((total_non_ascii_char_count/total_char_count) * 100, 2)
    else:
        total_char_to_non_ascii_char_ratio = 0
    # print(f"Non-ASCII characters in paragraphs: {non_ascii_char_count_paragraphs}")
    # # print(f"Total non-ASCII char count: {non_ascii_char_count}")
    # # print(f"Total char count: {char_count}")
    # print(f"Total char count: {total_char_count} to total non-ASCII char count: {total_non_ascii_char_count}. Ratio: {total_char_to_non_ascii_char_ratio} (%)")
    
    # Data export
    if not chosen_file:
        data_to_csv = [
            ["Data set", data_set],
            ["Document Name", "Total char count", "Total non-ASCII char count", "Total char to non-ASCII char ratio (%)"],
            [path.stem, total_char_count, total_non_ascii_char_count, total_char_to_non_ascii_char_ratio],
        ]
    else:
        data_to_csv = [
            ["Document Name", path.stem],
            ["Data set", data_set],
            ["Total char count", total_char_count],
            ["Total non-ASCII char count", total_non_ascii_char_count],
            ["Total char to non-ASCII char ratio (%)", total_char_to_non_ascii_char_ratio],
            ["Non-ASCII counts per paragraph", *non_ascii_char_count_paragraphs]
        ]
    
    return data_to_csv, 2
        
def main() -> None:
    unified_statistical_analysis_file.statistical_analysis('6_unicode', 'TEST_0', unicode_analysis)
    unified_statistical_analysis_file.statistical_analysis('6_unicode', None, unicode_analysis)

if __name__ == "__main__":
    main()