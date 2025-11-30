from pathlib import Path
from collections import Counter
import _1_element_count
from docx import Document
from docx.document import Document as DocumentObject
import unified_statistical_analysis_file

def number_of_runs_with_a_single_character(document: DocumentObject) -> int:
    count = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if len(run.text) == 1 and run.text != ' ':
                count += 1
    return count

def run_text_with_single_character(document: DocumentObject) -> list:
    runs = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if len(run.text) == 1 and run.text != ' ':
                runs.append(run.text)
    return runs

def bin_single_chars(single_char: str) -> str:
    if single_char.isupper():   
        return "uppercase"
    elif single_char.islower():
        return "lowercase"
    elif single_char.isdigit():
        return "digit"
    else:
        return "other"
    
def single_char_analysis(path: Path, data_set: str, chosen_file: bool) -> list[list]:
    document = Document(str(path))
    total_run_element_count = _1_element_count.count_total_runs_elements(document)
    single_char_run_count = number_of_runs_with_a_single_character(document)
    run_percentages = str(round((single_char_run_count / total_run_element_count) * 100, 2)) #.replace(".", ",")
    single_char_texts = run_text_with_single_character(document)
    single_char_texts_list = [bin_single_chars(single_char) for single_char in single_char_texts]
    single_char_bins = {key: Counter(single_char_texts_list).get(key, 0) for key in ['uppercase', 'lowercase', 'digit', 'other']}
    
    chars = []
    frequencies = []
    frequency_percentages_single_run = []
    frequency_percentages_total_runs = []

    # print(f"Number of runs with a single char (non-whitespace): {single_char_run_count} out of {total_run_element_count} runs. ({run_percentages}%).")
    # print("Char from each single-char run element:")
    # print(single_char_texts)
    for char, frequency in single_char_bins.items():
        if single_char_run_count != 0:
            frequency_percent_out_of_single_run = str(round((frequency / single_char_run_count) * 100, 2)) #.replace(".", ",")
        else:
            frequency_percent_out_of_single_run = 0
        if total_run_element_count != 0:
            frequency_percent_out_of_total_runs = str(round((frequency / total_run_element_count) * 100, 2)) #.replace(".", ",")
        else:
            frequency_percent_out_of_total_runs = 0
        chars.append(char)
        frequencies.append(frequency)
        frequency_percentages_single_run.append(frequency_percent_out_of_single_run)
        frequency_percentages_total_runs.append(frequency_percent_out_of_total_runs)
        # print(f"Char type: {char}. Frequency: {frequency} (Single char runs: {frequency_percent_out_of_single_run}%, Total runs: {frequency_percent_out_of_total_runs}%).")
    
    if not chosen_file:
        data_to_csv = [
            ["Document Name", path.stem],
            ["Data set", data_set],
            ["Single non-whitespace char run count", single_char_run_count],
            ["Total run element count", total_run_element_count],
            ["Total run element count percentage (%)", run_percentages],
            ["Char type", *chars],
            ["Char frequencies", *frequencies],
            ["Char frequencies (single run, %)", *frequency_percentages_single_run],
            ["Char frequencies (total runs, %)", *frequency_percentages_total_runs]
        ]
    else:
        data_to_csv = [
            ["Document Name", path.stem],
            ["Data set", data_set],
            ["Single non-whitespace char run count", single_char_run_count],
            ["Total run element count", total_run_element_count],
            ["Total run element count percentage (%)", run_percentages],
            ["Char type", *chars],
            ["Char frequencies", *frequencies],
            ["Char frequencies (single run, %)", *frequency_percentages_single_run],
            ["Char frequencies (total runs, %)", *frequency_percentages_total_runs]
        ]
    
    return data_to_csv

def main() -> None:
    unified_statistical_analysis_file.singular_check('4_single_char', 'TEST_0', single_char_analysis)
    unified_statistical_analysis_file.singular_check('4_single_char', None, single_char_analysis)

if __name__ == "__main__":
    main()