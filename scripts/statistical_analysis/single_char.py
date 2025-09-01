from docx import Document
from pathlib import Path

def count_text_elements(document) -> int:
    root = document.part.element
    docTextElement = root.xpath("//w:document/w:body/w:p/w:r/w:t")
    return len(docTextElement)

def number_of_runs_with_a_single_character(document) -> int:
    count = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if len(run.text) == 1:
                count += 1
    return count

def run_text_with_single_character(document) -> list:
    runs = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if len(run.text) == 1:
                if run.text != ' ':
                    runs.append(run.text)
    with open("results/one_char.txt", "w", encoding="utf-8") as file:
        for char in runs:
            file.write(str(char) + "\n")
    return runs
    

docPath_0 = Path("data_set/clean_files/TEST_0.docx")
docPath_1 = Path("data_set/stego-files/stego-method_1/TEST_0.docx")
docPath_2 = Path("data_set/stego-files/stego-method_4/TEST_0.docx")
docPath_3 = Path("data_set/stego-files/stego-method_5/TEST_0.docx")
docPath_4= Path("data_set/stego-files/stego-method_6/TEST_0.docx")

paths = [docPath_0, docPath_1, docPath_2, docPath_3, docPath_4]

for path in paths:
    print("")
    print(f"DOCUMENT: {path}")
    document = Document(path)
    print(f"Text element count: {count_text_elements(document)}")
    print(f"Number of runs with a single char: {number_of_runs_with_a_single_character(document)}")
    print("Char from each single-char run element:")
    print(run_text_with_single_character(document))
