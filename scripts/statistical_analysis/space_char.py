import re
from docx import Document
from pathlib import Path

def count_text_elements(document) -> int:
    root = document.part.element
    docTextElement = root.xpath("//w:document/w:body/w:p/w:r/w:t")
    return len(docTextElement)

def count_words_in_paragraphs(document) -> int:
    word_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
        words = re.findall(r'\S+', text, flags=re.UNICODE)
        word_count += len(words)
    return word_count

def count_whitespace_characters_in_each_run(document) -> list:
    whitespace_counts = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            whitespace_count = sum(1 for char in run.text if char.isspace())
            whitespace_counts.append(whitespace_count)
    with open("results/whitespace_counts.txt", "w", encoding="utf-8") as file:
        for count in whitespace_counts:
            file.write(str(count) + "\n")
    return whitespace_counts

docPath_0 = Path("data_set/clean_files/TEST_0.docx")
docPath_1 = Path("data_set/stego-files/stego-method_1/TEST_0.docx")
docPath_2 = Path("data_set/stego-files/stego-method_4/TEST_0.docx")
docPath_3 = Path("data_set/stego-files/stego-method_5/TEST_0.docx")
docPath_4= Path("data_set/stego-files/stego-method_6/TEST_0.docx")

paths = [docPath_0, docPath_1, docPath_2, docPath_3, docPath_4]

for path in paths:
    print("")
    print(f"DOCUMENT: {path}")
    document = Document(path)
    word_count = count_words_in_paragraphs(document)
    whitespace_count = count_whitespace_characters_in_each_run(document)
    whitespace_count_ = 0
    for count in whitespace_count:
        whitespace_count_ += count
    word_to_whitespace_ratio = 100 * (whitespace_count_/word_count)
    print(f"Whitespace count: {whitespace_count_}")
    print(f"Word count: {word_count}")
    print(f"Word count to whitespace ratio: {word_to_whitespace_ratio}")
