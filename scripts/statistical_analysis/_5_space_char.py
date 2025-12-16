import re
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentObject
import scripts.statistical_analysis.unified_statistical_analysis_file as unified_statistical_analysis_file

# Count the amount of words in each paragraph element
def count_words_in_each_paragraph(document: DocumentObject) -> int:
    word_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
        words = re.findall(r'\S+', text, flags=re.UNICODE)
        word_count += len(words)
    return word_count

# Count the amount of individual whitespace characters in each paragraph element
def count_whitespace_characters_in_each_paragraph(document: DocumentObject) -> list:
    whitespace_counts = []
    for paragraph in document.paragraphs:
        whitespace_count = sum(1 for char in paragraph.text if char.isspace())
        whitespace_counts.append(whitespace_count)
    return whitespace_counts

# Count the amount of individual whitespace characters in each run element
def count_whitespace_characters_in_each_run(document: DocumentObject) -> list:
    whitespace_counts = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            whitespace_count = sum(1 for char in run.text if char.isspace())
            whitespace_counts.append(whitespace_count)
    return whitespace_counts

### Main function ###
def space_char_analysis(path: Path, data_set: str, chosen_file: bool) -> tuple[list[list], int]:
    document = Document(str(path))
    total_word_count = count_words_in_each_paragraph(document)
    whitespace_counts_per_paragraph = count_whitespace_characters_in_each_paragraph(document)
    whitespace_counts_per_run = count_whitespace_characters_in_each_run(document)
    total_whitespace_count = 0
    for count in whitespace_counts_per_run:
        total_whitespace_count += count
    if total_word_count != 0:
        word_to_whitespace_ratio = round((total_whitespace_count/total_word_count) * 100, 2)
    else:
        word_to_whitespace_ratio = 0
    # print(f"Whitespace count in each paragraph: {whitespace_counts_per_paragraph}")
    # print(f"Whitespace count in each run: {whitespace_counts_per_run}")
    # print(f"Total whitespace count: {total_whitespace_count}")
    # print(f"Total word count: {total_word_count}")
    # print(f"Total word count: {total_word_count} to total whitespace count {total_whitespace_count}. Ratio: {word_to_whitespace_ratio} (%)")
    
    # Data export
    if not chosen_file:
        data_to_csv = [
            ["Data set", data_set],
            ["Document Name", "Total word count", "Total whitespace count", "Word to Whitespace ratio (%)"],
            [path.stem, total_word_count, total_whitespace_count, word_to_whitespace_ratio],
        ]
    else:
        data_to_csv = [
            ["Document Name", path.stem],
            ["Data set", data_set],
            ["Total word count", total_word_count],
            ["Total whitespace count", total_whitespace_count],
            ["Word to Whitespace ratio (%)", word_to_whitespace_ratio],
            ["Whitespace counts per paragraph", *whitespace_counts_per_paragraph] #,
            # ["Whitespace counts per run", *whitespace_counts_per_run]
        ]
    
    return data_to_csv, 2

def main() -> None:
    unified_statistical_analysis_file.statistical_analysis('5_space_char', 'TEST_0', space_char_analysis)
    unified_statistical_analysis_file.statistical_analysis('5_space_char', None, space_char_analysis)

if __name__ == "__main__":
    main()