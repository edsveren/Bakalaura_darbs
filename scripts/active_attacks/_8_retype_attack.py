from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.parser import OxmlElement
import scripts.active_attacks.unified_active_attack_file as unified_active_attack_file

def retype_attack(
        stegoDocPath: str, 
        attackedStegoDocPath: str
        ) -> None:
    
    document = Document(stegoDocPath)
    new_document = Document()

    # Copy stego-document content and add it to the new document content
    print(f"Copying all content of: {Path(stegoDocPath).name} into a new document")
    for paragraph in document.paragraphs:
        new_paragraph_element = OxmlElement('w:p')
        for run in paragraph.runs:
            current_run_element = run._r
            new_run_element = OxmlElement('w:r')
            for child_element in current_run_element:
                new_run_element.append(deepcopy(child_element))
            new_paragraph_element.append(new_run_element)
        new_document._element.body.append(new_paragraph_element)
    new_document.save(attackedStegoDocPath)

if __name__ == "__main__":
    unified_active_attack_file.unified_attack("08_retype_attack", retype_attack, False)