from docx import Document
from pathlib import Path
from docx.oxml.shared import OxmlElement
from copy import deepcopy

def retype_in_new_document(document, new_document):
    for paragraph in document.paragraphs:
        new_paragraph_element = OxmlElement('w:p')
        for run in paragraph.runs:
            current_run_element = run._r
            new_run_element = OxmlElement('w:r')
            for child_element in current_run_element:
                new_run_element.append(deepcopy(child_element))
            new_paragraph_element.append(new_run_element)
        new_document._element.body.append(new_paragraph_element)

base = "data_set/stego-files"
for directories in Path(base).iterdir():
    for file in directories.iterdir():
        docPath = f"{base}/{directories.name}/{file.name}" #Path("data_set/clean_files/TEST_0.docx")
        document = Document(docPath)
        new_document = Document()
        retype_in_new_document(document, new_document)
        stegoDocPath = Path(f"data_set/attacked_stego-files/8_retype_attack/{directories.name}/{file.name}")
        new_document.save(stegoDocPath)
        print("Saved:", stegoDocPath)