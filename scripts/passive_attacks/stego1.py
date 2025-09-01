import base64
from pathlib import Path
from docx import Document
from docx.oxml.shared import qn

# Extract text from the document
def extract_text(document) -> str:
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text.replace('\xa0', '\x20')) # NBSP -> space
    text = "\n".join(text)
    #print("Teksts no dokumenta:\n", text)
    return text

# Stego-message
def stego_message() -> list[str]:
    stegoMessageText = Path("stego-messages/stego-message.txt").read_text(encoding="utf-8")
    #print("Stego-message data:", stegoMessageText)
    return stegoMessageText

# Stego-message in Base64
def stego_message_base64(stegoMessage_bytes) -> tuple[str, bytes]:
    stegoMessage_toBase64_bytes = base64.b64encode(stegoMessage_bytes)
    stegoMessage_toBase64_text = stegoMessage_toBase64_bytes.decode('ascii')
    #print("Stego-message Base64:", stegoMessage_toBase64_text)
    return stegoMessage_toBase64_text, stegoMessage_toBase64_bytes

# Extraction algorithm
def stego_message_extraction(document) -> str:
    stegoMessage_as_base64 = ''
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run_properties = run._r.rPr
            if run_properties != None:
                if len(run.text) == 1:
                    color_element = run_properties.find(qn('w:color'))
                    font_size_element = run_properties.find(qn('w:sz'))
                    vanish_element = run_properties.find(qn('w:vanish'))

                    if None not in (color_element, font_size_element, vanish_element):
                        color_element_value = color_element.get(qn('w:val'))
                        font_size_value = font_size_element.get(qn('w:val'))
                        #vanish_element_value = vanish_element.get(qn('w:val'))
                        if color_element_value.upper() == 'FFFFFF' and font_size_value == '2':
                            stegoMessage_as_base64 += run.text
    #print(stegoMessage_as_base64)
    stegoMessage = base64.b64decode(stegoMessage_as_base64).decode('utf-8')
    #print(stegoMessage)
    return stegoMessage
            
# DOCX file
docPath = Path("data_set/attacked_stego-files/stego-method_1/TEST_0.docx")
document = Document(docPath)
text = extract_text(document)

stego_message_text = stego_message()

### Main

print("Extracting stego-message...")
stego_message_extracted = stego_message_extraction(document)
if stego_message_extracted != '':
    print(f"The extracted stego-message: {stego_message_extracted}")
    if stego_message_text == stego_message_extracted:
        print("Extraction successful!")
    else:
        print("Extracted message is not equal to stego-message!")
else:
    print(f"The extracted stego-message: {None}")
