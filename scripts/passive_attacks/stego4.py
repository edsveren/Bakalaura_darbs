from pathlib import Path
from docx import Document
from docx.oxml.shared import qn

# Extract text from the document
def extract_text(document) -> str:
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text.replace('\xa0', '\x20')) # NBSP -> space
    text = "\n".join(text)
    #print("Teksts no dokumenta:\n", text)
    return text

# Stego-message
def stego_message() -> list[str]:
    stegoMessageText = Path("stego-messages/stego-message.txt").read_text(encoding="utf-8")
    #print("Stego-message data:", stegoMessageText)
    return stegoMessageText

# Extraction algorithm
def stego_message_extraction(document) -> str:
    stegoMessage_bytes = b''
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run_properties = run._r.rPr
            if run_properties != None:
                color_element = run_properties.find(qn('w:color'))
                if color_element != None:
                    color_element_value = color_element.get(qn('w:val'))
                    if color_element_value.lower() not in ('000000', 'auto'):
                        r_bit = int(color_element_value[0:2], 16)
                        g_bit = int(color_element_value[2:4], 16)
                        b_bit = int(color_element_value[4:6], 16)
                        if r_bit <= 7 and g_bit <= 7 and b_bit <= 3:
                            stego_bit_string = f"{r_bit:03b}{g_bit:03b}{b_bit:02b}"
                            stego_byte = int(stego_bit_string, 2).to_bytes(1, 'big')
                            #print(stego_byte)
                            stegoMessage_bytes += stego_byte
    #print(stegoMessage_bytes)
    stegoMessage = stegoMessage_bytes.decode('utf-8')
    #print(stegoMessage)
    return stegoMessage
            
# DOCX file
base = "data_set/attacked_stego-files/stego-method_4"
base = "data_set/stego-files/stego-method_4"
for file in Path(base).iterdir():
    docPath = f"{base}/{file.name}" #Path("data_set/clean_files/TEST_0.docx")
    #docPath = Path("data_set/attacked_stego-files/stego-method_1/TEST_0.docx")
    document = Document(docPath)
    text = extract_text(document)

    stego_message_text = stego_message()

    # Main

    print("Extracting stego-message...")
    stego_message_extracted = stego_message_extraction(document)
    if stego_message_extracted != '':
        print(f"The extracted stego-message: {stego_message_extracted}")
        if stego_message_text == stego_message_extracted:
            print("Extraction successful!")
        else:
            print("Extracted message is not equal to stego-message!")
    else:
        print(f"The extracted stego-message: {None}")
    print()