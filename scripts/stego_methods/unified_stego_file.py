import random
from pathlib import Path
from typing import Callable, Any
from docx import Document
from docx.text.run import Run
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn

# Get stego-message text and bytes from file
def stego_message() -> tuple[str, bytes]:
    stego_message_text = Path("stego_messages/stego_message.txt").read_text(encoding="utf-8")
    stego_message_bytes = stego_message_text.encode("utf-8")
    return stego_message_text, stego_message_bytes

# Extract text from the document
def extract_text(
        document: DocumentObject,
        NBSP: bool,
        index: int
    ) -> str:
    text = []
    for paragraph in document.paragraphs[index:]:
        if NBSP == True:
            text.append(paragraph.text.replace('\xa0', '\x20'))
        else:
            text.append(paragraph.text)
    text = ''.join(text)
    # print(f"Document text:\n{text}")
    return text

# Choose a random paragraph to embed the stego-message into
def choose_random_paragraph(
        document: DocumentObject,
        necessary_element_count_function: Callable[[DocumentObject, int], int],
        is_capacity_enough_for_message: Callable[[DocumentObject, int, int, int], bool],
        stego_message_size_bits: int
    ) -> int | None:

    paragraphs = document.paragraphs
    if not paragraphs:
        return None
    
    # Keep looking for a valid randomly generated paragraph
    # Until a paragraph with enough capacity is found
    while True:
        random_paragraph_index = random.randint(0, len(paragraphs) - 1)
        random_paragraph = paragraphs[random_paragraph_index]
        necessary_element_count = necessary_element_count_function(document, random_paragraph_index)
        is_valid = is_capacity_enough_for_message(document, necessary_element_count, stego_message_size_bits, random_paragraph_index)
        if is_valid:
            print(f"Random paragraph start: {random_paragraph.text}")
            return random_paragraph_index

# Get stego-message bytes and bits     
def get_bytes_and_bits(stego_message_bytes: bytes|str) -> tuple[int, int]:
    stego_message_size_bytes = len(stego_message_bytes)
    stego_message_size_bits = 8 * stego_message_size_bytes
    
    # print(f"Stego-message bytes: {stego_message_size_bytes}")
    # print(f"Stego-message bites: {stego_message_size_bits}")

    return stego_message_size_bytes, stego_message_size_bits

### Main file ###   
def stego_method(
        stego_method: str, # 'stego_method_1', 'stego_method_2', etc.
        specialized_stego_message: tuple[str, bytes]|None, # For specialized stego-message input
        necessary_element_count_function: Callable[[DocumentObject, int], int], # Count words, chars, black chars, etc.
        is_capacity_enough_for_message:  Callable[[DocumentObject, int, int, int], bool], # Check document capacity function
        embedding_algorithm: Callable[[Run, Any, int, int], int], # Embedding algorithm function, Any = str|bytes
        embedding_data_type: str, # 'string' or 'bytes'
        extraction_algorithm: Callable[[DocumentObject], str] # Extract stego-message function
    ) -> None:
    base = "data_set/clean_files"
    # Loop through all DOCX files in the clean directory
    for file in Path(base).iterdir():
        docPath = str(Path(f"{base}/{file.name}"))
        print(f"DOCX file: {docPath}")
        print("Beginning the embedding process...")
        document = Document(docPath)
        text = extract_text(document, True, 0)

        # Count the full corresponding document elements (words, chars, black chars, etc.)
        necessary_element_count = necessary_element_count_function(document, 0)

        # Get stego-message text and bytes from default file or specialized input
        if specialized_stego_message == None:
            stego_message_text, stego_message_bytes = stego_message()
        else:
            stego_message_text, stego_message_bytes = specialized_stego_message
        
        # Get stego-message size in bytes and bits
        # Unispace stego-method adds padding to the stego-message
        # So it increases the embedded data
        if stego_method == 'stego_method_5':
            stego_message_size_bytes, stego_message_size_bits = get_bytes_and_bits(stego_message_text)
        else:
            stego_message_size_bytes, stego_message_size_bits = get_bytes_and_bits(stego_message_bytes)

        embedded = False
        while not embedded:
            print("Checking if the cover object is valid for embedding...")
            is_valid = is_capacity_enough_for_message(document, necessary_element_count, stego_message_size_bits, 0)
            print("The cover object is valid:", is_valid)

            if not is_valid:
                print("Not enough capacity in the document to embed the message.")
                break

            print("Embedding stego-message...")
            random_paragraph_index = choose_random_paragraph(document, necessary_element_count_function, is_capacity_enough_for_message, stego_message_size_bits)
            # This should not happen since the capacity was already determined to be enough but just in case
            if random_paragraph_index == None:
                print("No paragraphs available for embedding.")
                break
            
            # Initialize the starting index and payload for embedding
            # Unicode homoglyphs stego-method adds a '1' at the start of the stego-message bits
            # To allow for proper extraction, so the payload needs to be increased by 1 bit
            if stego_method == 'stego_method_6':
                payload = stego_message_size_bits + 1
            else:
                payload = stego_message_size_bytes
            stego_index = 0

            # Determine the data to embed based on the specified type (most use string)
            data_to_embed = ''
            if embedding_data_type == 'bytes':
                data_to_embed = stego_message_bytes
            elif embedding_data_type == 'string':
                data_to_embed = stego_message_text

            for paragraph in document.paragraphs [random_paragraph_index:]:
                if stego_index < payload:
                    # Get the original amount of runs in the paragraph
                    # Because runs will be added or otherwise modified during embedding
                    # Which leads to endless loops or skipped runs
                    original_run_amount = list(paragraph.runs)
                    for run in original_run_amount:
                        run_element = run._r
                        # Only process runs that contain text elements
                        if run_element.find(qn('w:t')) != None:
                            if stego_index < payload:
                                # The main embedding function call
                                # It works within each individual run
                                # To embed as much of the stego-message as possible in that run
                                # And returns the updated stego index after embedding 
                                stego_index = embedding_algorithm(run, data_to_embed, stego_index, payload)
                            else:
                                break
                else:
                    break
            
            # When getting the original stego-message for comparison
            # Unispace method is different
            # It compromises the original stego-message for embedding purposes
            # First off, it can only use uppercase ASCII characters
            # So all lowercase ASCII characters are converted to uppercase 
            # While non-ASCII characters are removed
            # And secondly, it uses whitespace padding characters at the start and end
            # To denote the start and end of the stego-message for the purposes of extraction
            if stego_method == 'stego_method_5':
                stego_message_text = stego_message_text[1:-1]
            else:
                stego_message_text, _ = stego_message()

            extracted_stego_message = extraction_algorithm(document)
            print(f"Stego-message: {stego_message_text}")
            print(f"Extracted stego-message: {extracted_stego_message}")
            
            # Compare the extracted stego-message with the one before embedding
            # If they don't match, something went wrong with the embedding process
            # That's not an intended behavior
            if stego_message_text != extracted_stego_message:
                print("Extracted message is not equal to stego-message!")
                break
            print("Embedding successful!")
            embedded = True

        if embedded:
            stegoDocPath = str(Path(f"data_set/stego_files/{stego_method}/{file.name}"))
            document.save(stegoDocPath)
            print(f"Saved: {stegoDocPath}")

            if file.name == 'TEST_0.docx':
                stegoDocPath_TEST_0 = str(Path(f"data_set/TEST_0/{stego_method}_{file.name}"))
                document.save(stegoDocPath_TEST_0)
                print(f"Saved: {stegoDocPath_TEST_0}")
        else:
            print("Embedding not possible.")
        print()