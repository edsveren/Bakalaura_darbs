import re
import random
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.run import Run
from docx.oxml.shared import OxmlElement, qn

# Count chars in paragraphs
def count_chars_in_paragraphs(document, index) -> int:
    char_count = 0
    for paragraph in document.paragraphs[index:]:
        text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
        chars = re.findall(r'\S', text, flags=re.UNICODE)
        char_count += len(chars)
    #print("Kopējais vārdu skaits:", char_count)
    return char_count

# Check if max capacity is enough for the message
def is_capacity_enough_for_message(char_count, stegoMessage_size_bits) -> bool:
    cap = 8 * char_count
    is_valid = stegoMessage_size_bits <= cap
    return is_valid

# Extract text from the document
def extract_text(document) -> str:
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text.replace('\xa0', '\x20')) # NBSP -> space
    text = "\n".join(text)
    text = ''.join(text)
    #print("Teksts no dokumenta:\n", text)
    return text

# Stego-message
def stego_message() -> tuple[list[str], bytes]:
    stegoMessageText = Path("stego-messages\stego-message.txt").read_text(encoding="utf-8")
    stegoMessage_bytes = stegoMessageText.encode("utf-8")
    #print("Stego-message data:", stegoMessageText)
    return stegoMessageText, stegoMessage_bytes

# Choose random paragraph
def choose_random_paragraph(document, stegoMessage_size_bits) -> int | None:
    paragraphs = document.paragraphs
    if not paragraphs:
        return None
    while True:
        random_paragraph_index = random.randint(0, len(paragraphs) - 1)
        random_paragraph = paragraphs[random_paragraph_index]
        char_count = count_chars_in_paragraphs(document, random_paragraph_index)
        is_valid = is_capacity_enough_for_message(char_count, stegoMessage_size_bits)
        if is_valid:
            #print("Random paragraph start:", random_paragraph.text)
            return random_paragraph_index

# Create a new run
def insert_in_run(previous_run, char, stego_bits, base_run) -> Run|None:
    current_run_element = previous_run._r
    base_run_element = base_run._r

    # New run element
    new_run_element = OxmlElement('w:r')

    # New run properties
    if base_run_element.rPr is not None:
        # Copy all existing run properties
        new_run_element.append(deepcopy(base_run_element.rPr))
        run_properties = new_run_element.find(qn('w:rPr'))
    else:
        # Create an empty run properties element
        run_properties = OxmlElement('w:rPr')
        new_run_element.insert(0, run_properties)

    # New text element
    text_element = OxmlElement('w:t')
    new_run_element.append(text_element)

    text_element.text = char
    
    if text_element.text.startswith('\x20') or text_element.text.endswith('\x20'):
        text_element.set(qn('xml:space'), 'preserve')

    if stego_bits != None:
        color_element = run_properties.find(qn('w:color'))
        if color_element is None:
            color_element = OxmlElement('w:color')
            run_properties.append(color_element)
        color_element_value = color_element.get(qn('w:val')) 
        if color_element_value == None or color_element_value.lower() in ('auto', '000000'):
            r_bit = int(stego_bits[0:3], 2)
            g_bit = int(stego_bits[3:6], 2)
            b_bit = int(stego_bits[6:8], 2)

            rgb_string = f"{r_bit:02x}{g_bit:02x}{b_bit:02x}"
            color_element.set(qn('w:val'), rgb_string)
            #run_properties.append(color_element)
        else:
            return None
    
    current_run_element.addnext(new_run_element)
    new_run = Run(new_run_element, previous_run._parent)
    return new_run

# Splitting the existing runs into before, current and after
def slipt_run_for_embedding(run, byte) -> Run | None:
    text = run.text
    cover_chars = re.search("\S", text, flags=re.UNICODE)
    if cover_chars is None:
        return None

    cover_symbol_index = cover_chars.start()
    single_cover_char = text[cover_symbol_index]
    left_text = text[:cover_symbol_index] # text before the first whitespace
    right_text = text[cover_symbol_index + 1:] # text after the first whitespace

    stego_byte_to_binary_string = f"{byte:08b}"

    stego_char_run = insert_in_run(run, single_cover_char, stego_byte_to_binary_string, run)
    if stego_char_run == None:
        return None
    
    # left text
    run.text = left_text

    # right text
    remaining_run = insert_in_run(stego_char_run, right_text, None, run)
    return remaining_run

# Embedding algorithm
def embedding_in_run(run, stegoMessage_bytes, stego_index, payload) -> int:
    current_run = run
    #text = run.text.replace('\xa0', '\x20') # NBSP -> space
    non_whitespace_chars = re.findall(r'\S', run.text, flags=re.UNICODE)
    nr_of_non_whitespace_chars = len(non_whitespace_chars)

    for _ in range(nr_of_non_whitespace_chars):
        if stego_index < payload:
            next_run = slipt_run_for_embedding(current_run, stegoMessage_bytes[stego_index])
            if next_run != None:
                current_run = next_run
                stego_index += 1
            else:
                break
        else:
            break
    return stego_index

# Extraction algorithm
def stego_message_extraction(document) -> str:
    stegoMessage_bytes = b''
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run_properties = run._r.rPr
            if run_properties != None:
                color_element = run_properties.find(qn('w:color'))
                if color_element != None:
                    color_element_value = color_element.get(qn('w:val'))
                    if color_element_value.lower() not in ('000000', 'auto'):
                        r_bit = int(color_element_value[0:2], 16)
                        g_bit = int(color_element_value[2:4], 16)
                        b_bit = int(color_element_value[4:6], 16)
                        if r_bit <= 7 and g_bit <= 7 and b_bit <= 3:
                            stego_bit_string = f"{r_bit:03b}{g_bit:03b}{b_bit:02b}"
                            stego_byte = int(stego_bit_string, 2).to_bytes(1, 'big')
                            #print(stego_byte)
                            stegoMessage_bytes += stego_byte
    #print(stegoMessage_bytes)
    stegoMessage = stegoMessage_bytes.decode('utf-8')
    #print(stegoMessage)
    return stegoMessage

### Main ###    
# DOCX file
base = "data_set/clean_files"
for file in Path(base).iterdir():
    docPath = f"{base}/{file.name}" #Path("data_set/clean_files/TEST_0.docx")
    #docPath = Path("data_set/clean_files/TEST_0.docx")
    print(docPath)
    document = Document(docPath)
    text = extract_text(document)
    char_count = count_chars_in_paragraphs(document, 0)

    stego_message_text, stegoMessage_bytes = stego_message()
    stegoMessage_size_bytes = len(stegoMessage_bytes)
    stegoMessage_size_bits = 8 * stegoMessage_size_bytes
    #print("Regular bytes:", stegoMessage_size_bytes)
    #print("Regular bites:", stegoMessage_size_bits)

    embedded = False
    while not embedded:
        # Check if the paragraph has enough runs to embed the message
        is_valid = is_capacity_enough_for_message(char_count, stegoMessage_size_bits)
        print("The cover object is valid:", is_valid)
        if not is_valid:
            print("Not enough capacity in the document to embed the message.")
            break

        random_paragraph_index = choose_random_paragraph(document, stegoMessage_size_bits)
        if random_paragraph_index is None:
            print("No paragraphs available for embedding.")
            break

        # Embed stego-message in DOCX
        print("Embedding stego-message...")
        payload = stegoMessage_size_bytes
        stego_index = 0
        #while payload < stego_index:
        for paragraph in document.paragraphs [random_paragraph_index:]:
                if stego_index < payload:
                    original_run_amount = list(paragraph.runs)
                    for run in original_run_amount:
                        run_element = run._r
                        if run_element.find(qn('w:t')) != None:
                            if stego_index < payload:
                                stego_index = embedding_in_run(run, stegoMessage_bytes, stego_index, payload)
                            else:
                                break
                else:
                    break

        #print("Extracting stego-message...")
        extracted_stego_message = stego_message_extraction(document)
        if stego_message_text != extracted_stego_message:
            print(extracted_stego_message)
            print("Extracted message is not equal to stego-message!")
            break
        #print("Extraction successful!")
        print("Embedding successful!")     
        embedded = True

    if embedded:
        stegoDocPath = Path(f"data_set/stego-files/stego-method_4/{file.name}")
        document.save(stegoDocPath)
        print("Saved:", stegoDocPath)
    else:
        print("Embedding not possible.")
    print()