import re
import nltk
from nltk.corpus import treebank
import random
from pathlib import Path
from copy import deepcopy
from collections import Counter
from docx import Document
from docx.text.run import Run
from docx.oxml.shared import OxmlElement, qn
# def formula(tagged):
#     derigie_tag = ['NN', 'IN', 'DT', 'JJ', 'NNS', 'RB', 'VB', 'PRP', 'VBG', 'VBZ']
#     for i in range(0, 11):
#         for j in range(0, 11):
#             if tagged.tag in derigie_tag:
#                 tagged.word = 
#         if tag.startswith('NN'):
#     return value
# cap2 = formula

# Count words in paragraphs
def count_words_in_paragraphs(doc, index) -> int:
    word_count = 0
    for paragraph in doc.paragraphs[index:]:
        text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
        words = re.findall(r'\S+', text, flags=re.UNICODE)
        word_count += len(words)
    #print("Kopējais vārdu skaits:", word_count)
    return word_count

# Check if max capacity is enough for the message
def is_capacity_enough_for_message(layers, stegoMessage_size_bits) -> bool:
    # Formula = 
    cap = 0
    for layer in layers:
        for word in layer:
           cap += len(word) * 12
    is_valid = stegoMessage_size_bits <= cap
    return is_valid

# Extract text from the document
def extract_text(doc) -> str:
    text = []
    for p in doc.paragraphs:
        text.append(p.text.replace('\xa0', '\x20')) # NBSP -> space
    text = "\n".join(text)
    #print("Teksts no dokumenta:\n", text)
    return text

# Stego-message
def stego_message() -> tuple[list[str], bytes]:
    stegoMessagePath = Path("stego_message.txt")
    stegoMessage_bytes = stegoMessagePath.read_bytes()
    stegoMessageText = []
    with open(stegoMessagePath, encoding="utf-8") as input:
        for line in input:
            stegoMessageText.append(line)
    #print("Stego-message data:", stegoMessageText)
    return stegoMessageText, stegoMessage_bytes

# Convert stego-message to binary
def stego_message_to_binary(stegoMessage_bytes):
    stegoMessage_toBinary = ''.join(format(byte, '08b') for byte in stegoMessage_bytes)
    print("TODO")
    

# Choose layers based on POS tags
def choose_layers() -> list[str]:
    layers = []
    for tag, count in tag_counts.items():
        if count > 1:
            layers.append(tag)
    return layers

# Choose random paragraph
def choose_random_paragraph(doc) -> int | None:
    paragraphs = doc.paragraphs
    if not paragraphs:
        return None
    while True:
        random_paragraph_index = random.randint(0, len(paragraphs) - 1)
        random_paragraph = paragraphs[random_paragraph_index]
        word_count = count_words_in_paragraphs(doc, random_paragraph_index)
        is_valid = is_capacity_enough_for_message(word_count, stegoMessage_toBase64_size_bits)
        if is_valid:
            print("Random paragraph start:", random_paragraph.text)
            return random_paragraph_index

def insert_in_run(previous_run, char, type, base_run) -> Run:
    current_run_element = previous_run._r
    base_run_element = base_run._r

    # New run element
    new_run_element = OxmlElement('w:r')

    # New run properties
    if base_run_element.rPr is not None:
        # Copy all existing run properties
        new_run_element.append(deepcopy(base_run_element.rPr))
        run_properties = new_run_element.find(qn('w:rPr'))
    else:
        # Create an empty run properties element
        run_properties = OxmlElement('w:rPr')
        new_run_element.insert(0, run_properties)

    # New text element
    text_element = OxmlElement('w:t')
    new_run_element.append(text_element)

    # Insert character based on type
    match type:
        case 'whitespace':
            text_element.text = char
            #if text_element.text.startswith('\x20') or text_element.text.endswith('\x20'):
            text_element.set(qn('xml:space'), 'preserve')
        case 'stego_char': # len(char) == 1 and char != ('\x20', '\xa0')
            text_element.text = char
            color_element = OxmlElement('w:color')
            color_element.set(qn('w:val'), 'FFFFFF')
            run_properties.append(color_element)
        case _:
            text_element.text = char           
    
    current_run_element.addnext(new_run_element)
    return Run(new_run_element, previous_run._parent)

def slipt_run_for_embedding(run, char) -> Run | None:
    text = run.text
    whitespace = text.find('\x20')
    if whitespace == -1:
        return None

    left_text = text[:whitespace] # text before the first whitespace
    right_text = text[whitespace + 1:] # text after the first whitespace

    # left text
    run.text = left_text

    # left whitespace
    left_whitespace = insert_in_run(run, '\x20', 'whitespace', run)

    # stego character
    stego_char = insert_in_run(left_whitespace, char, 'stego_char', run)

    # right whitespace
    right_whitespace = insert_in_run(stego_char, '\x20', 'whitespace', run)

    # right text
    remaining_run = insert_in_run(right_whitespace, right_text, None, run)
    return remaining_run

def embedding_in_run(run, stegoMessage_toBase64_text, stego_index, payload) -> int:
    current_run = run
    #text = run.text.replace('\xa0', '\x20') # NBSP -> space
    text_whitespaces = re.findall(r'\x20', run.text, flags=re.UNICODE)
    nr_of_unused_whitespace = len(text_whitespaces)

    for _ in range(nr_of_unused_whitespace):
        if stego_index < payload:
            next_run = slipt_run_for_embedding(current_run, stegoMessage_toBase64_text[stego_index])
            if next_run != None:
                current_run = next_run
                stego_index += 1
            else:
                break
        else:
            break
    return stego_index
            
# DOCX file
docPath = Path("TEST_CASES\TEST_0\TEST_0.docx")
doc = Document(docPath)
text = extract_text(doc)
tokens = nltk.word_tokenize(text)
print("Vārdu skaits:", len(tokens))
word_count = count_words_in_paragraphs(doc, 0)
tagged = nltk.pos_tag(tokens)

stego_message_text, stegoMessage_bytes = stego_message()
stegoMessage_size_bytes = len(stegoMessage_bytes)
stegoMessage_size_bits = 8 * stegoMessage_size_bytes
#print("Regular bytes:", stegoMessage_size_bytes)
#print("Regular bites:", stegoMessage_size_bits)

# 2
# need english text
#print("Vārdu skaits ar POS tagiem:", len(tagged))
amount = 0
#for word, tag in tagged:
#        print(f"{word}  ->  {tag}")
tag_counts = Counter(tag for _, tag in tagged)
print(tag_counts) 
unique_tag_count = len(tag_counts)
print("Unikālo POS tagu skaits:", unique_tag_count)



# Main

embedded = False
while not embedded:
    # Check if the paragraph has enough runs to embed the message
    is_valid = is_capacity_enough_for_message(word_count, stegoMessage_toBase64_size_bits)
    print("The cover object is valid:", is_valid)
    if not is_valid:
        print("Not enough capacity in the document to embed the message.")
        break

    random_paragraph_index = choose_random_paragraph(doc)
    if random_paragraph_index is None:
        print("No paragraphs available for embedding.")
        break

    # Embed stego-message in DOCX
    print("Embedding stego-message...")
    payload = len(stegoMessage_toBase64_text)
    stego_index = 0
    #while payload < stego_index:
    for paragraph in doc.paragraphs [random_paragraph_index:]:
            if stego_index < payload:
                original_run_amount = list(paragraph.runs)
                for run in original_run_amount:
                    if stego_index < payload:
                        stego_index = embedding_in_run(run, stegoMessage_toBase64_text, stego_index, payload)
                    else:
                        break
            else:
                break

    embedded = True
    print("Embedding successful!")


if embedded:
    stegoDocPath = docPath.with_name(docPath.stem + "_STEGO.docx")
    doc.save(stegoDocPath)
    print("Saved:", stegoDocPath)
else:
    print("Embedding not possible.")





# docPath = Path("TEST_CASES\TEST_0\TEST_0.docx")
# doc = Document(docPath)
# text = extract_text(doc)
# tokens = nltk.word_tokenize(text)
# print("Vārdu skaits:", len(tokens))
# #word_count = count_words_in_paragraphs(doc, 0)
# tagged = nltk.pos_tag(tokens)

# tag_counts = Counter(tag for _, tag in tagged)
# print(tag_counts)                 # piem., {'NN': 42, 'DT': 15, ...}
# print(tag_counts.most_common(10)) # TOP10

# total_tag_occurrences = sum(tag_counts.values())
# print("Kopējais tagu skaits:", total_tag_occurrences)

# unique_tag_count = len(tag_counts)
# print("Unikālo tagu skaits:", unique_tag_count)

# x_values = list(tag_counts.values())
# print("Tagu skaita vērtības:", x_values)

# layers = tag_counts.most_common(10)
# print(tag_counts.most_common(10))

# # i need list(tag_counts.keys()) for layers
# # each element in layers is made up of tag_counts.values()
# layers = [tag for tag, _ in layers]