import re
import random
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.run import Run
from docx.oxml.shared import OxmlElement, qn

# S
StegoState = {
    "rgb_marker_state": [1, 2, 2],   # [r, g, b] ; r ∈ {-1,1}, g/b ∈ {-1,1,2}
    "bit_delta_state": [0, 0],           # [g_delta, b_delta] ; 0..7
}

def reset_state():
    r_bit_marker = 1 # -1 | 1 (differential marker)
    g_bit_marker = 2 # -1 | 1 | 2 (backward direction, forward direction, start/end marker)
    b_bit_marker = 2 # -1 | 1 | 2 (backward direction, forward direction, start/end marker)
    StegoState["rgb_marker_state"] = [r_bit_marker, g_bit_marker, b_bit_marker]

    g_bit_delta = 0
    b_bit_delta = 0
    StegoState["bit_delta_state"] = [g_bit_delta, b_bit_delta]

# Count (black coloured) chars in paragraphs starting from index
def count_black_chars_in_paragraphs(document: Document, index: int) -> int:
    black_char_count = 0
    for paragraph in document.paragraphs[index:]:
        for run in paragraph.runs:
            text = ''
            run_properties = run._r.rPr
            if run_properties == None:
                text = run.text.replace('\xa0', '\x20')  # NBSP -> space
            else:
                color_element = run_properties.find(qn('w:color'))
                if color_element == None:
                    text = run.text.replace('\xa0', '\x20')  # NBSP -> space
                else:
                    color_element_value = color_element.get(qn('w:val')) 
                    if color_element_value == None or color_element_value.lower() in ('auto', '000000'):
                        text = run.text.replace('\xa0', '\x20')  # NBSP -> space
                    else:
                        text = ''
            #text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
            chars = re.findall(r'\S', text, flags=re.UNICODE)
            black_char_count += len(chars)
    #print("Kopējais vārdu skaits:", char_count)
    return black_char_count

# Check if max capacity is enough for the message
def is_capacity_enough_for_message(black_char_count: int, stegoMessage_size_bits: int) -> bool:
    # in theory
    # cap = 2 * (char_count - 1)
    # in practice
    cap = 2 * (black_char_count - 1) + 4
    is_valid = stegoMessage_size_bits <= cap
    return is_valid

# Extract text from the document
def extract_text(document: Document) -> str:
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text.replace('\xa0', '\x20')) # NBSP -> space
    text = "\n".join(text)
    text = ''.join(text)
    #print("Teksts no dokumenta:\n", text)
    return text

# Stego-message
def stego_message() -> tuple[list[str], bytes]:
    stegoMessageText = Path("stego_messages\stego_message.txt").read_text(encoding="utf-8")
    stegoMessage_bytes = stegoMessageText.encode("utf-8")
    #print("Stego-message data:", stegoMessageText)
    return stegoMessageText, stegoMessage_bytes

# Choose random paragraph
def choose_random_paragraph(document: Document, stegoMessage_size_bits: int) -> int | None:
    paragraphs = document.paragraphs
    if not paragraphs:
        return None
    while True:
        random_paragraph_index = random.randint(0, len(paragraphs) - 1)
        random_paragraph = paragraphs[random_paragraph_index]
        black_char_count = count_black_chars_in_paragraphs(document, random_paragraph_index)
        is_valid = is_capacity_enough_for_message(black_char_count, stegoMessage_size_bits)
        if is_valid:
            #print("Random paragraph start:", random_paragraph.text)
            return random_paragraph_index

# Insert a bit in color channel
def insert_bit(stego_bit: str, color_channel_index: int, bit_delta: int) -> int:
    bit_direction = StegoState["rgb_marker_state"][color_channel_index + 1]
    #bit_delta = StegoState["bit_delta_state"][color_channel_index]

    match stego_bit:
        case '2':
            bit_delta = 2
            StegoState["rgb_marker_state"][color_channel_index + 1] = 1
            StegoState["bit_delta_state"][color_channel_index] = bit_delta
            return bit_delta
        case '1':
            if bit_direction == 1:
                if bit_delta < 7:
                    bit_delta += 1
                    if bit_delta == 7:
                        StegoState["rgb_marker_state"][color_channel_index + 1] = -1
                else:
                    StegoState["rgb_marker_state"][color_channel_index + 1] = -1
            elif bit_direction == -1:
                if bit_delta > 0:
                    bit_delta -= 1
                    if bit_delta == 0:
                        StegoState["rgb_marker_state"][color_channel_index + 1] = 1
                else:
                    StegoState["rgb_marker_state"][color_channel_index + 1] = 1

            StegoState["bit_delta_state"][color_channel_index] = bit_delta
            return bit_delta
        case _:  # '0'
            return bit_delta

# Create a new run
def insert_in_run(previous_run: Run, char: str, stego_bits: str, base_run: Run) -> Run|None:
    current_run_element = previous_run._r
    base_run_element = base_run._r

    # New run element
    new_run_element = OxmlElement('w:r')

    # New run properties
    if base_run_element.rPr is not None:
        # Copy all existing run properties
        new_run_element.append(deepcopy(base_run_element.rPr))
        run_properties = new_run_element.find(qn('w:rPr'))
    else:
        # Create an empty run properties element
        run_properties = OxmlElement('w:rPr')
        new_run_element.insert(0, run_properties)

    # New text element
    text_element = OxmlElement('w:t')
    new_run_element.append(text_element)

    text_element.text = char
    
    if text_element.text.startswith('\x20') or text_element.text.endswith('\x20'):
        text_element.set(qn('xml:space'), 'preserve')

    if stego_bits != None:
        #print("Stego bits to insert:", stego_bits)
        color_element = run_properties.find(qn('w:color'))
        if color_element is None:
            color_element = OxmlElement('w:color')
            run_properties.append(color_element)
        color_element_value = color_element.get(qn('w:val')) 
        if color_element_value == None or color_element_value.lower() in ('auto', '000000'):
            r_bit = 0 # int(color_element_value[0:2], 16)
            g_bit = 0 # int(color_element_value[2:4], 16)
            b_bit = 0 # int(color_element_value[4:6], 16)

            r_bit_marker = StegoState["rgb_marker_state"][0]
            g_bit_marker = StegoState["rgb_marker_state"][1]
            b_bit_marker = StegoState["rgb_marker_state"][2]

            g_bit_delta = StegoState["bit_delta_state"][0]
            b_bit_delta = StegoState["bit_delta_state"][1]
                            
            if g_bit_marker != 2 and b_bit_marker != 2 and stego_bits != '2':
                stego_bit_0, stego_bit_1 = stego_bits

                g_bit = insert_bit(stego_bit_0, 0, g_bit_delta)
                b_bit = insert_bit(stego_bit_1, 1, b_bit_delta)
                if g_bit == g_bit_delta and b_bit == b_bit_delta:
                # if g_bit == b_bit:
                    # r_bit = r_bit + r_bit_marker
                    # if r_bit_marker == 0 and r_bit == 0:
                    #     r_bit = r_bit + 1
                    #     StegoState["rgb_marker_state"][0] = 1
                    #     #rgb_marker_direction[0] = -1
                    # elif r_bit_marker == 1: #if r_bit_marker == -1
                    #     StegoState["rgb_marker_state"][0] = 0
                    #     #rgb_marker_direction[0] = 1
                    # if r_bit_marker == 1:
                    #     StegoState["rgb_marker_state"][0] = 0
                    #     #rgb_marker_direction[0] = -1
                    # elif r_bit_marker == 0: #if r_bit_marker == -1
                    #     StegoState["rgb_marker_state"][0] = 1
                    #     #rgb_marker_direction[0] = 1
                    StegoState["rgb_marker_state"][0] ^= 1
                    r_bit_marker = StegoState["rgb_marker_state"][0]
                    r_bit = r_bit_marker
                else:
                    r_bit = r_bit_marker
                    #r_bit = 0
            else:
                g_bit = insert_bit(stego_bits, 0, g_bit_delta)
                b_bit = insert_bit(stego_bits, 1, b_bit_delta)
                StegoState["rgb_marker_state"][0] = 1

            rgb_string = f"{r_bit:02x}{g_bit:02x}{b_bit:02x}"
            # if r_bit < 0 or g_bit < 0 or b_bit < 0:
            #     print(f"rgb: {rgb_string}")
            # elif r_bit > 255 or g_bit > 255 or b_bit > 255:
            #     print(f"rgb: {rgb_string}")
            #print(f"rgb: {rgb_string}")
            
            color_element.set(qn('w:val'), rgb_string)
            #run_properties.append(color_element)
        else:
            return None
    
    current_run_element.addnext(new_run_element)
    new_run = Run(new_run_element, previous_run._parent)
    return new_run

# Splitting the existing runs into before, current and after
def slipt_run_for_embedding(run: Run, bit_pair: str) -> Run|None:
    text = run.text
    cover_chars = re.search("\S", text, flags=re.UNICODE)
    if cover_chars is None:
        return None

    cover_symbol_index = cover_chars.start()
    single_cover_char = text[cover_symbol_index]
    left_text = text[:cover_symbol_index] # text before the first whitespace
    right_text = text[cover_symbol_index + 1:] # text after the first whitespace

    #stego_byte_to_binary_string = f"{bit_pair:08b}"

    stego_char_run = insert_in_run(run, single_cover_char, bit_pair, run)
    if stego_char_run == None:
        return None
    else:
        # left text
        if left_text != '': # Check if there is text to the left
            text_element = run._r.find(qn('w:t'))
            if text_element == None:
                text_element = OxmlElement('w:t')
                run._r.append(text_element)
            if left_text.startswith(' ') or left_text.endswith(' '):
                text_element.set(qn('xml:space'), 'preserve')
            run.text = left_text

        # right text
        remaining_run = insert_in_run(stego_char_run, right_text, None, run)
        # Clean up
        if left_text == '':
            left_parent = run._r.getparent()
            left_parent.remove(run._r)
        return remaining_run

# Embedding algorithm
def embedding_in_run(run: Run, stegoMessage_toBinary: str, stego_index: int, payload: int) -> int:
    current_run = run
    #text = run.text.replace('\xa0', '\x20') # NBSP -> space
    non_whitespace_chars = re.findall(r'\S', run.text, flags=re.UNICODE)
    nr_of_non_whitespace_chars = len(non_whitespace_chars)

    for _ in range(nr_of_non_whitespace_chars):
        if stego_index < payload:
            if StegoState["rgb_marker_state"][1] != 2 and StegoState["rgb_marker_state"][2] != 2:
                bit_pair = stegoMessage_toBinary[stego_index:stego_index + 2]
                next_run = slipt_run_for_embedding(current_run, bit_pair)
                if next_run != None:
                    current_run = next_run
                    stego_index += 2
                    if stego_index == payload:
                        StegoState["rgb_marker_state"][1] = 2
                        StegoState["rgb_marker_state"][2] = 2
                else:
                    break
            else:
                next_run = slipt_run_for_embedding(current_run, '2')
                if next_run != None:
                    current_run = next_run
                continue
        else:
            next_run = slipt_run_for_embedding(current_run, '2')
            break
    return stego_index

# Extraction algorithm
def stego_message_extraction(document: Document, stegoMessage_toBinary) -> str:
    bytes_num = 0
    stegoMessage_bytes = b''
    stego_message_state = 0
    stego_bit_string = ''
    stego_bits_size = 0
    Stego_Bit_String_For_Checking = []
    try:
        for paragraph in document.paragraphs:    
            if stego_message_state == 2:
                break
            for run in paragraph.runs:
                g_bit_prev, b_bit_prev = StegoState["bit_delta_state"]
                run_properties = run._r.rPr
                if run_properties != None:
                    color_element = run_properties.find(qn('w:color'))
                    if color_element != None:
                        color_element_value = color_element.get(qn('w:val'))
                        if color_element_value.lower() not in ('000000', 'auto'):
                            g_bit = int(color_element_value[2:4], 16)
                            b_bit = int(color_element_value[4:6], 16)
                            if g_bit <= 7 and b_bit <= 7:
                                match stego_message_state:
                                    case 0:
                                        # stego-message not found yet
                                        if g_bit == 2 and b_bit == 2:
                                            StegoState["bit_delta_state"][0] = g_bit
                                            StegoState["bit_delta_state"][1] = b_bit
                                            stego_message_state += 1
                                    case 1:
                                        # stego-message found, extracting
                                        g_bit_delta = abs(g_bit - g_bit_prev)
                                        b_bit_delta = abs(b_bit - b_bit_prev)
                                        
                                        if g_bit_delta == 2 and b_bit_delta == 2:
                                            stego_message_state += 1
                                            break
                                        else:
                                            stego_bit_string += f"{g_bit_delta}{b_bit_delta}"
                                            stego_bits_size += 2
                                            # print(f"index: {stego_bits_size}, {stego_bits_size + 1}")
                                            # print(f"deltas: {g_bit_delta, b_bit_delta}")
                                            # print(f"from: {g_bit_prev}, {b_bit_prev}")
                                            # print(f"to: {g_bit}, {b_bit}")
                                            StegoState["bit_delta_state"][0] = g_bit
                                            StegoState["bit_delta_state"][1] = b_bit
                                            Stego_Bit_String_For_Checking.append(f"{g_bit_delta}{b_bit_delta}")
                                            if len(stego_bit_string) == 8:
                                                stego_byte = int(stego_bit_string, 2).to_bytes(1, 'big')
                                                #print(stego_byte)
                                                stegoMessage_bytes += stego_byte
                                                stego_bit_string = ''
                                                bytes_num += 1
                                    case 2:
                                        # stego-message end found, stop extracting
                                        break
        #print(stegoMessage_bytes)
        stegoMessage = stegoMessage_bytes.decode('utf-8')
        #print(stegoMessage)
        return stegoMessage
    except:
        stegoMessage_toBinary_wrong = ''.join(Stego_Bit_String_For_Checking)
        for i in range(len(stegoMessage_toBinary)):
            # if stegoMessage_toBinary[i] == stegoMessage_toBinary_wrong[i]:
            #     print(f"Same at index {i}: {stegoMessage_toBinary[i]}")
            if stegoMessage_toBinary[i] != stegoMessage_toBinary_wrong[i]:
                print(f"Different at index {i}: should be {stegoMessage_toBinary[i]}, is {stegoMessage_toBinary_wrong[i]}")

### Main ###             
# DOCX file
base = "data_set/clean_files"
#for file in Path(base).iterdir():
#docPath = f"{base}/{file.name}"
docPath = Path("data_set/clean_files/TEST_0.docx")
print(f"DOCX file: {docPath}")
document = Document(docPath)
text = extract_text(document)
black_char_count = count_black_chars_in_paragraphs(document, 0)

stego_message_text, stegoMessage_bytes = stego_message()
stegoMessage_size_bytes = len(stegoMessage_bytes)
stegoMessage_size_bits = 8 * stegoMessage_size_bytes
#print("Regular bytes:", stegoMessage_size_bytes)
#print("Regular bites:", stegoMessage_size_bits)
stegoMessage_toBinary = ''.join(format(byte, '08b') for byte in stegoMessage_bytes)
#print("Stego-message in binary:", stegoMessage_toBinary)

reset_state()
embedded = False
while not embedded:
    # Check if the paragraph has enough runs to embed the message
    print("Checking if the cover object is valid for embedding...")
    is_valid = is_capacity_enough_for_message(black_char_count, stegoMessage_size_bits)
    print("The cover object is valid:", is_valid)
    if not is_valid:
        print("Not enough capacity in the document to embed the message.")
        break

    random_paragraph_index = choose_random_paragraph(document, stegoMessage_size_bits)
    if random_paragraph_index is None:
        print("No paragraphs available for embedding.")
        break

    # Embed stego-message in DOCX
    print("Embedding stego-message...")
    payload = stegoMessage_size_bits #stegoMessage_size_bytes
    stego_index = 0
    for paragraph in document.paragraphs [random_paragraph_index:]:
        if stego_index < payload:
            original_run_amount = list(paragraph.runs)
            for run in original_run_amount:
                run_element = run._r
                # Only process runs that contain text
                if run_element.find(qn('w:t')) != None:
                    if stego_index < payload:
                        stego_index = embedding_in_run(run, stegoMessage_toBinary, stego_index, payload)
                    else:
                        break
        else:
            break
    print("Embedding successful!")
    reset_state()    
    #print("Extracting stego-message...")
    extracted_stego_message = stego_message_extraction(document, stegoMessage_toBinary)
    if stego_message_text != extracted_stego_message:
        print(extracted_stego_message)
        print("Extracted message is not equal to stego-message!")
        break
    #print("Extraction successful!")     
    embedded = True

if embedded:
    stegoDocPath = Path(f"data_set/stego_files/stego_method_3/TEST_0.docx") #{file.name}
    document.save(stegoDocPath)
    print("Saved:", stegoDocPath)
else:
    print("Embedding not possible.")
print()