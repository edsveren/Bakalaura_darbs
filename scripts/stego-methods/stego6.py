import re
import random
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.run import Run
from docx.oxml.shared import OxmlElement, qn

unicode_dictionary = {
    'a': '\u0430',
    'b': '\u042C',
    'c': '\u03F2',
    'd': '\u0501',
    'e': '\u0435',
    'f': '\uAB35',
    'g': '\u0261',
    'h': '\u04BB',
    'i': '\u0456',
    'j': '\u03F3',
    'k': '\u043A',
    'l': '\u04CF',
    'm': '\u043C',
    'n': '\u0578',
    'o': '\u03BF',
    'p': '\u0440',
    'q': '\u051B',
    'r': '\u1D26',
    's': '\u0455',
    't': '\u03C4',
    'u': '\u057D',
    'v': '\u1D20',
    'w': '\u051D',
    'x': '\u0445',
    'y': '\u0443',
    'z': '\u1D22'
}

reverse_unicode_dictionary = {value: key for key, value in unicode_dictionary.items()}

# Count words in paragraphs
def count_chars_in_paragraphs(document, index) -> int:
    char_count = 0
    for paragraph in document.paragraphs[index:]:
        text = paragraph.text.replace('\xa0', '\x20')  # NBSP -> space
        chars = re.findall(r'[a-z]', text, flags=re.UNICODE)
        char_count += len(chars)
    #print("Kopējais burtu skaits:", char_count)
    return char_count

# Check if max capacity is enough for the message
def is_capacity_enough_for_message(char_count, stegoMessage_size_bits) -> bool:
    cap = char_count
    is_valid = stegoMessage_size_bits <= cap
    return is_valid

# Extract text from the document
def extract_text(document) -> str:
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text.replace('\xa0', '\x20')) # NBSP -> space
    text = "\n".join(text)
    text = ''.join(text)
    #print("Teksts no dokumenta:\n", text)
    return text

# Stego-message
def stego_message() -> tuple[list[str], bytes]:
    stegoMessageText = Path("stego-messages\stego-message.txt").read_text(encoding="utf-8")
    stegoMessage_bytes = stegoMessageText.encode("utf-8")
    #print("Stego-message data:", stegoMessageText)
    return stegoMessageText, stegoMessage_bytes

# Transform stego-message to bit string
def stego_message_to_bit_string(stegoMessage_bytes) -> str:
    stego_byte_to_binary_string = ''
    for byte in stegoMessage_bytes:
        stego_byte_to_binary_string += f"{byte:08b}"
    #print(len(stego_byte_to_binary_string))
    return stego_byte_to_binary_string

# Choose random paragraph
def choose_random_paragraph(document, stegoMessage_toBase64_size_bits) -> int | None:
    paragraphs = document.paragraphs
    if not paragraphs:
        return None
    while True:
        random_paragraph_index = random.randint(0, len(paragraphs) - 1)
        random_paragraph = paragraphs[random_paragraph_index]
        word_count = count_chars_in_paragraphs(document, random_paragraph_index)
        is_valid = is_capacity_enough_for_message(word_count, stegoMessage_toBase64_size_bits)
        if is_valid:
            print("Random paragraph start:", random_paragraph.text)
            return random_paragraph_index

# Embedding algorithm
def embedding_in_run(run, stego_message_text, stego_index, payload) -> int:
    run_element = run._r
    run_properties = run_element.rPr
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_element.insert(0, run_properties)
    if run_properties.find(qn("w:noProof")) is None:
        run_properties.append(OxmlElement("w:noProof"))

    all_lowercase_letters = re.findall(r'[a-z]', run.text, flags=re.UNICODE)
    nr_of_unused_lowercase_letters = len(all_lowercase_letters)

    text = run.text
    embedded_text = ''
    remaining_text = text

    for _ in range(nr_of_unused_lowercase_letters):
        if stego_index >= payload:
            break

        lowercase_letter = re.search("[a-z]", remaining_text, flags=re.UNICODE)
        if lowercase_letter is None:
            embedded_text += remaining_text
            remaining_text = ''
            break
        else:
            cover_symbol_index = lowercase_letter.start()
            single_cover_char = remaining_text[cover_symbol_index]
            embedded_text += remaining_text[:cover_symbol_index]
            remaining_text = remaining_text[cover_symbol_index + 1:]
            if stego_message_text[stego_index] == '1':
                homoglyph = unicode_dictionary.get(single_cover_char)
                embedded_text += homoglyph
            else:
                embedded_text += single_cover_char
            stego_index += 1

    if remaining_text != '':
        embedded_text += remaining_text

    if embedded_text != text:
        run.text = embedded_text

    return stego_index

# Extraction algorithm
def stego_message_extraction(document) -> str:
    text = extract_text(document)
    stegoMessage_bytes = b''
    first_stego_char_found = False
    stego_byte_string = ''
    end_byte_string = '00000000'
    for char in text:
        stego_char = reverse_unicode_dictionary.get(char)
        if not first_stego_char_found:
            if stego_char != None:
                first_stego_char_found = True
            continue
        else:
            if stego_char != None:
                stego_byte_string += '1'
            elif 'a' <= char <= 'z':
                stego_byte_string += '0'
            else:
                continue
            if len(stego_byte_string) == 8:
                if stego_byte_string == end_byte_string:
                    break
                else:
                    stego_byte = int(stego_byte_string, 2).to_bytes(1, 'big')
                    stegoMessage_bytes += stego_byte
                    stego_byte_string = ''
    stegoMessage = stegoMessage_bytes.decode('utf-8')
    #print(stegoMessage)
    return stegoMessage

# DOCX file
docPath = Path("data_set/clean_files/TEST_0.docx")
document = Document(docPath)
text = extract_text(document)
word_count = count_chars_in_paragraphs(document, 0)

stego_message_text, stegoMessage_bytes = stego_message()
stegoMessage_size_bytes = len(stegoMessage_bytes)
stegoMessage_size_bits = 8 * stegoMessage_size_bytes
#print("Regular bytes:", stegoMessage_size_bytes)
#print("Regular bites:", stegoMessage_size_bits)

stegoMessage_bytes_to_binary_string = stego_message_to_bit_string(stegoMessage_bytes)
stegoMessage_bytes_to_binary_string = '1' + stegoMessage_bytes_to_binary_string

### Main

embedded = False
while not embedded:
    # Check if the paragraph has enough runs to embed the message
    is_valid = is_capacity_enough_for_message(word_count, stegoMessage_size_bits)
    print("The cover object is valid:", is_valid)
    if not is_valid:
        print("Not enough capacity in the document to embed the message.")
        break

    random_paragraph_index = choose_random_paragraph(document, stegoMessage_size_bits)
    if random_paragraph_index is None:
        print("No paragraphs available for embedding.")
        break

    # Embed stego-message in DOCX
    print("Embedding stego-message...")
    payload = stegoMessage_size_bits + 1
    stego_index = 0
    #while payload < stego_index:
    for paragraph in document.paragraphs [random_paragraph_index:]:
            if stego_index >= payload:
                break               
            else:
                #original_run_amount = list(paragraph.runs)
                for run in paragraph.runs:
                    run_element = run._r
                    if run_element.find(qn('w:t')) != None:
                        if stego_index >= payload:
                            break
                        else:
                            stego_index = embedding_in_run(run, stegoMessage_bytes_to_binary_string, stego_index, payload)

    print("Embedding successful!")

    print("Extracting stego-message...")
    if stego_message_text != stego_message_extraction(document):
        print("Extracted message is not equal to stego-message!")
        break
    print("Extraction successful!")     
    embedded = True

if embedded:
    stegoDocPath = Path(f"data_set/stego-files/stego-method_6/{docPath.name}")
    document.save(stegoDocPath)
    print("Saved:", stegoDocPath)
else:
    print("Embedding not possible.")